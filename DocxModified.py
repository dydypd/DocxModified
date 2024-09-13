import os
import re
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document

regex_notes_file = "regex_notes.txt"


def load_file():
    global file_path
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if file_path:
        document = Document(file_path)
        content = ""
        for para in document.paragraphs:
            content += para.text + "\n"
        original_text_box.delete("1.0", tk.END)
        original_text_box.insert(tk.END, content)
        preview_text_box.delete("1.0", tk.END)


def remove_characters():
    phrase_to_remove = entry.get()
    content = original_text_box.get("1.0", tk.END)

    if newline_var.get() == 1:
        content = content.replace("\n", " ")

    if regex_var.get() == 1:
        try:
            content = re.sub(phrase_to_remove, "", content)
        except re.error:
            messagebox.showerror("Lỗi", "Biểu thức chính quy không hợp lệ!")
            return
    else:
        content = content.replace(phrase_to_remove, "")

    preview_text_box.delete("1.0", tk.END)
    preview_text_box.insert(tk.END, content)


def save_file():
    if not file_path:
        messagebox.showerror("Lỗi", "Chưa có tệp mà đòi sửa :)))")
        return
    file_directory, file_name = os.path.split(file_path)
    file_name_no_ext, file_ext = os.path.splitext(file_name)
    new_file_name = file_name_no_ext + "_Edited" + file_ext
    new_file_path = os.path.join(file_directory, new_file_name)

    document = Document()
    modified_content = preview_text_box.get("1.0", tk.END)

    for line in modified_content.splitlines():
        document.add_paragraph(line)
    document.save(new_file_path)
    messagebox.showinfo("Thành công", f"Tệp đã được lưu thành cmn công dưới tên {new_file_name}!")


def save_regex_note():
    regex_note = regex_entry.get()
    if regex_note.strip() == "":
        messagebox.showerror("Lỗi", "Biểu thức chính quy không được để trống!")
        return

    with open(regex_notes_file, "a") as f:
        f.write(regex_note + "\n")

    regex_listbox.insert(tk.END, regex_note)
    regex_entry.delete(0, tk.END)


def load_regex_notes():
    if os.path.exists(regex_notes_file):
        with open(regex_notes_file, "r") as f:
            lines = f.readlines()
            for line in lines:
                regex_listbox.insert(tk.END, line.strip())


def use_selected_regex(event):
    selected_regex = regex_listbox.get(regex_listbox.curselection())
    entry.delete(0, tk.END)
    entry.insert(0, selected_regex)


root = tk.Tk()
root.title("Xóa Ký Tự trong Tệp DOCX với Preview hehe")

file_path = ""

btn_load = tk.Button(root, text="Tải Tệp", command=load_file)
btn_load.pack(pady=5)

entry = tk.Entry(root)
entry.pack(pady=5)

newline_var = tk.IntVar()
newline_checkbox = tk.Checkbutton(root, text="Xoá dấu xuống dòng", variable=newline_var)
newline_checkbox.pack(pady=5)
regex_var = tk.IntVar()
regex_checkbox = tk.Checkbutton(root, text="Sử dụng Biểu thức chính quy", variable=regex_var)
regex_checkbox.pack(pady=5)
btn_remove = tk.Button(root, text="Xoá Ký Tự và Xem Trước", command=remove_characters)
btn_remove.pack(pady=5)

tk.Label(root, text="Ghi chú công thức RegEx:").pack()
regex_entry = tk.Entry(root)
regex_entry.pack(pady=5)

btn_save_regex = tk.Button(root, text="Lưu Công Thức RegEx", command=save_regex_note)
btn_save_regex.pack(pady=5)

tk.Label(root, text="Các công thức RegEx đã lưu:").pack()
regex_listbox = tk.Listbox(root, height=5, width=80)
regex_listbox.pack(pady=5)

regex_listbox.bind('<<ListboxSelect>>', use_selected_regex)

load_regex_notes()

tk.Label(root, text="Nội dung gốc:").pack()
original_text_box = tk.Text(root, height=10, width=80)
original_text_box.pack(pady=5)

tk.Label(root, text="Nội dung sau khi thay đổi (Xem trước):").pack()
preview_text_box = tk.Text(root, height=10, width=80)
preview_text_box.pack(pady=5)
btn_save = tk.Button(root, text="Lưu Tệp", command=save_file)
btn_save.pack(pady=5)

root.mainloop()
